#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created on Sun Nov 18 14:52:36 2018

@author: doan
"""

from docx import Document
doc = Document()
doc.add_paragraph('Namelist parameter')



f = open("../namelist.ucm")
lines = f.readlines()
f.close()

table = doc.add_table(rows=len(lines)+1, cols=(3))
table.style = 'Table Grid'
table.cell(0,0).text = 'Parameter'
table.cell(0,1).text = 'Example'
table.cell(0,2).text = 'Description'


for i,l in enumerate(lines): 
    table.cell(1+i,0).text = l.split("=")[0]
    if len(l.split("=")) > 1:  
       table.cell(1+i,1).text = l.split("=")[1]
    #if v[2]: 
    #    #print v[1]
    #    table.cell(1+i, 1).merge(table.cell(1+i, max_dom)).text = v[1]
    #else: 
    #    for j in range(max_dom): table.cell(1+i, j+1).text = v[1][j]



doc.save('namelist.docx')


